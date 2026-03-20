#!/usr/bin/env python3
"""
make_bilingual_doc.py — Генератор билингвального Word-документа DE+RU
v2.0: поддержка мульти-частей (619-1, 619-2, 619-3) с заголовками секций
			поддержка форматов "HH:MM:SS: текст" и "HH:MM:SS текст"
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════════════
# КОНСТАНТЫ
# ═══════════════════════════════════════════════════════════════════════════

# Матчит "HH:MM:SS:" или "HH:MM:SS " в начале строки
LINE_RE     = re.compile(r'^(\d{2}:\d{2}:\d{2}):?\s+(.*)', re.DOTALL)
# Заголовок секции: "619-1 (июль 1942)" или "619-2 (июль 1942)" и т.д.
SECTION_RE  = re.compile(r'^(\d{3,}-\d+\s*\([^)]+\))\s*$')
SPEAKER_RE  = re.compile(r'^[^:]{1,30}:\s+')
TS_INLINE   = re.compile(r'\b(\d{2}:\d{2}:\d{2})\b')

COL_WIDTHS_CM    = [2.8, 7.6, 7.6]
MATCH_TOLERANCE  = 5  # секунд

COLOR_HEADER_BG  = "1F3864"
COLOR_HEADER_FG  = "FFFFFF"
COLOR_SECTION_BG = "2E4057"   # чуть светлее шапки — строка-разделитель части
COLOR_SECTION_FG = "FFFFFF"
COLOR_TS_BG      = "D9E1F2"
COLOR_DE_BG      = "FFFFFF"
COLOR_RU_BG      = "F2F2F2"
COLOR_EMPTY      = "BFBFBF"

FONT_MAIN  = "Times New Roman"
FONT_TS    = "Courier New"
SIZE_HDR   = 11
SIZE_BODY  = 10
SIZE_TS    = 10


# ═══════════════════════════════════════════════════════════════════════════
# ПАРСИНГ
# ═══════════════════════════════════════════════════════════════════════════

def ts_to_sec(ts: str) -> int:
		h, m, s = ts.split(":")
		return int(h) * 3600 + int(m) * 60 + int(s)


def parse_txt(path: Path) -> list[dict]:
		"""
		Парсит TXT с поддержкой:
		- Заголовков секций:  "619-1 (июль 1942)"
		- Таймкодов:          "00:00:02: текст"  или  "00:00:02 текст"
		- Меток спикеров:     "Диктор: текст"  (удаляются)
		- Разделителей ====   (игнорируются)

		Возвращает список блоков:
			{'section': str, 'ts': str, 'sec': int, 'text': str}
		"""
		blocks  = []
		current = None
		section = ""

		for raw in path.read_text(encoding="utf-8").splitlines():
				line = raw.strip()
				if not line:
						continue
				if set(line) <= set("=- \t"):
						continue

				# Заголовок секции?
				ms = SECTION_RE.match(line)
				if ms:
						section = ms.group(1).strip()
						continue

				# Строка с таймкодом?
				m = LINE_RE.match(line)
				if m:
						ts   = m.group(1)
						text = m.group(2).strip()
						# Убираем "Диктор: " и подобные метки
						spk = SPEAKER_RE.match(text)
						if spk:
								text = text[spk.end():]
						if current:
								blocks.append(current)
						current = {
								"section": section,
								"ts":      ts,
								"sec":     ts_to_sec(ts),
								"text":    text
						}
				else:
						if current:
								current["text"] += " " + line

		if current:
				blocks.append(current)

		return blocks


def split_by_inner_timestamps(blocks: list[dict]) -> list[dict]:
		"""Разбивает блоки с внутренними таймкодами на суб-блоки."""
		result = []
		for block in blocks:
				text  = block["text"]
				parts = TS_INLINE.split(text)
				if len(parts) == 1:
						result.append(block)
						continue
				pre = parts[0].strip()
				if pre:
						result.append({**block, "text": pre})
				i = 1
				while i < len(parts) - 1:
						inner_ts   = parts[i]
						inner_text = parts[i + 1].strip()
						if inner_text:
								result.append({
										"section": block["section"],
										"ts":      inner_ts,
										"sec":     ts_to_sec(inner_ts),
										"text":    inner_text
								})
						i += 2
		return result


# ═══════════════════════════════════════════════════════════════════════════
# ВЫРАВНИВАНИЕ DE ↔ RU  (с учётом секций)
# ═══════════════════════════════════════════════════════════════════════════

def align_blocks(de_blocks: list[dict], ru_blocks: list[dict]) -> list[dict]:
		"""
		Выравнивает DE и RU блоки по (section, sec).
		Возвращает строки таблицы:
			{'section': str, 'ts': str, 'sec': int, 'de': str, 'ru': str}
		"""
		# Группируем по секциям
		de_by_section: dict[str, list] = {}
		for b in de_blocks:
				de_by_section.setdefault(b["section"], []).append(b)

		ru_by_section: dict[str, list] = {}
		for b in ru_blocks:
				ru_by_section.setdefault(b["section"], []).append(b)

		all_sections = list(dict.fromkeys(
				[b["section"] for b in de_blocks] +
				[b["section"] for b in ru_blocks]
		))

		rows = []

		for section in all_sections:
				de_sec = de_by_section.get(section, [])
				ru_sec = ru_by_section.get(section, [])

				# Маркер начала секции (заголовок-разделитель)
				rows.append({
						"section":    section,
						"ts":         "",
						"sec":        -1,
						"de":         "",
						"ru":         "",
						"is_section": True
				})

				all_secs = sorted(set(
						[b["sec"] for b in de_sec] +
						[b["sec"] for b in ru_sec]
				))

				ru_used = set()
				de_used = set()
				ru_idx  = {b["sec"]: b for b in ru_sec}
				de_idx  = {b["sec"]: b for b in de_sec}

				for sec in all_secs:
						# Найти DE
						de_block = None
						if sec in de_idx and id(de_idx[sec]) not in de_used:
								de_block = de_idx[sec]
						else:
								for b in de_sec:
										if id(b) not in de_used and abs(b["sec"] - sec) <= MATCH_TOLERANCE:
												de_block = b
												break

						# Найти RU
						ru_block = None
						if sec in ru_idx and id(ru_idx[sec]) not in ru_used:
								ru_block = ru_idx[sec]
						else:
								for b in ru_sec:
										if id(b) not in ru_used and abs(b["sec"] - sec) <= MATCH_TOLERANCE:
												ru_block = b
												break

						if de_block is None and ru_block is None:
								continue

						ts = de_block["ts"] if de_block else ru_block["ts"]
						rows.append({
								"section":    section,
								"ts":         ts,
								"sec":        sec,
								"de":         de_block["text"] if de_block else "",
								"ru":         ru_block["text"] if ru_block else "",
								"is_section": False
						})

						if de_block:
								de_used.add(id(de_block))
						if ru_block:
								ru_used.add(id(ru_block))

		return rows


# ═══════════════════════════════════════════════════════════════════════════
# WORD HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
		tc   = cell._tc
		tcPr = tc.get_or_add_tcPr()
		shd  = OxmlElement("w:shd")
		shd.set(qn("w:val"),   "clear")
		shd.set(qn("w:color"), "auto")
		shd.set(qn("w:fill"),  hex_color)
		tcPr.append(shd)


def _set_col_width(cell, width_cm: float):
		tc   = cell._tc
		tcPr = tc.get_or_add_tcPr()
		tcW  = OxmlElement("w:tcW")
		tcW.set(qn("w:w"),    str(int(width_cm * 567)))
		tcW.set(qn("w:type"), "dxa")
		tcPr.append(tcW)


def _add_paragraph(cell, text: str, font_name: str, font_size: int,
									 bold=False, color_hex: str = None,
									 align=WD_ALIGN_PARAGRAPH.LEFT):
		para = cell.paragraphs[0]
		para.alignment = align
		run = para.add_run(text)
		run.font.name = font_name
		run.font.size = Pt(font_size)
		run.font.bold = bold
		if color_hex:
				r = int(color_hex[0:2], 16)
				g = int(color_hex[2:4], 16)
				b = int(color_hex[4:6], 16)
				run.font.color.rgb = RGBColor(r, g, b)
		return para


def _merge_row_cells(row, bg_color: str, text: str):
		"""Объединяет все 3 ячейки строки в одну (для заголовка секции)."""
		cells = row.cells
		cells[0].merge(cells[1])
		cells[0].merge(cells[2])
		_set_cell_bg(cells[0], bg_color)
		_add_paragraph(cells[0], text,
									 font_name=FONT_MAIN, font_size=SIZE_HDR,
									 bold=True, color_hex=COLOR_SECTION_FG,
									 align=WD_ALIGN_PARAGRAPH.CENTER)


# ═══════════════════════════════════════════════════════════════════════════
# ГЕНЕРАЦИЯ ДОКУМЕНТА
# ═══════════════════════════════════════════════════════════════════════════

def build_doc(rows: list[dict], doc_title: str) -> Document:
		doc = Document()

		for section in doc.sections:
				section.top_margin    = Cm(1.5)
				section.bottom_margin = Cm(1.5)
				section.left_margin   = Cm(1.5)
				section.right_margin  = Cm(1.5)

		# Заголовок
		tp = doc.add_paragraph()
		tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
		tr = tp.add_run(doc_title)
		tr.font.name = FONT_MAIN
		tr.font.size = Pt(14)
		tr.font.bold = True
		doc.add_paragraph()

		# Таблица
		table = doc.add_table(rows=1, cols=3)
		table.style = "Table Grid"

		# Шапка
		hdr = table.rows[0].cells
		for i, (cell, label) in enumerate(zip(hdr, ["Таймкод", "Немецкий", "Русский"])):
				_set_cell_bg(cell, COLOR_HEADER_BG)
				_add_paragraph(cell, label, FONT_MAIN, SIZE_HDR,
											 bold=True, color_hex=COLOR_HEADER_FG,
											 align=WD_ALIGN_PARAGRAPH.CENTER)
				_set_col_width(cell, COL_WIDTHS_CM[i])

		# Строки
		for row in rows:
				tr_cells = table.add_row().cells

				if row.get("is_section"):
						# Строка-разделитель секции (объединённая ячейка)
						_merge_row_cells(table.rows[-1], COLOR_SECTION_BG,
														 f"▶  {row['section']}")
						continue

				# Таймкод
				_set_cell_bg(tr_cells[0], COLOR_TS_BG)
				_add_paragraph(tr_cells[0], row["ts"],
											 FONT_TS, SIZE_TS, bold=True,
											 align=WD_ALIGN_PARAGRAPH.CENTER)
				_set_col_width(tr_cells[0], COL_WIDTHS_CM[0])

				# DE
				_set_cell_bg(tr_cells[1], COLOR_DE_BG)
				if row["de"]:
						_add_paragraph(tr_cells[1], row["de"], FONT_MAIN, SIZE_BODY)
				else:
						_add_paragraph(tr_cells[1], "—", FONT_MAIN, SIZE_BODY,
													 color_hex=COLOR_EMPTY,
													 align=WD_ALIGN_PARAGRAPH.CENTER)
				_set_col_width(tr_cells[1], COL_WIDTHS_CM[1])

				# RU
				_set_cell_bg(tr_cells[2], COLOR_RU_BG)
				if row["ru"]:
						_add_paragraph(tr_cells[2], row["ru"], FONT_MAIN, SIZE_BODY)
				else:
						_add_paragraph(tr_cells[2], "—", FONT_MAIN, SIZE_BODY,
													 color_hex=COLOR_EMPTY,
													 align=WD_ALIGN_PARAGRAPH.CENTER)
				_set_col_width(tr_cells[2], COL_WIDTHS_CM[2])

		return doc


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("  make_bilingual_doc.py v2.1 — DE+RU билингвальная таблица")
    print("=" * 60)

    folder_input = input("\n📂 Путь к папке с файлами: ").strip().replace('"', '')
    folder = Path(folder_input)
    if not folder.exists() or not folder.is_dir():
        print(f"❌ Папка не найдена: {folder}")
        sys.exit(1)

    # Ищем DE и RU файлы по суффиксу
    de_candidates = list(folder.glob("*de.txt")) + list(folder.glob("*DE.txt"))
    ru_candidates = list(folder.glob("*ru.txt")) + list(folder.glob("*RU.txt"))

    if not de_candidates:
        print("❌ Не найден файл *de.txt в папке")
        sys.exit(1)

    de_path = de_candidates[0]
    ru_path = ru_candidates[0] if ru_candidates else None

    print(f"   ✅ DE: {de_path.name}")
    print(f"   {'✅' if ru_path else '⚠️ '} RU: {ru_path.name if ru_path else 'не найден — колонка будет пустой'}")

    # Заголовок документа
    stem = de_path.stem
    stem = re.sub(r'-?de$|-?ru$', '', stem, flags=re.IGNORECASE)
    stem = stem.replace("-", " ")
    doc_title = f"Перевод DDW {stem}"

    # Выходной файл — в ту же папку
    out_path = folder / (de_path.stem.replace("de", "").strip("-_ ") + ".docx")

    print(f"\n📖 Парсинг DE: {de_path.name}")
    de_blocks = split_by_inner_timestamps(parse_txt(de_path))
    sections  = list(dict.fromkeys(b["section"] for b in de_blocks))
    print(f"   → {len(de_blocks)} блоков | секции: {sections}")

    ru_blocks = []
    if ru_path:
        print(f"📖 Парсинг RU: {ru_path.name}")
        ru_blocks = split_by_inner_timestamps(parse_txt(ru_path))
        print(f"   → {len(ru_blocks)} блоков")
    else:
        print("ℹ️  RU не указан — колонка будет пустой")

    print(f"\n🔗 Выравнивание...")
    rows = align_blocks(de_blocks, ru_blocks)
    data_rows = [r for r in rows if not r.get("is_section")]
    print(f"   → {len(data_rows)} строк данных в {len(sections)} секциях")

    print(f"\n📝 Генерация DOCX...")
    doc = build_doc(rows, doc_title)
    doc.save(str(out_path))

    print(f"\n✅ Готово: {out_path}")


if __name__ == "__main__":
		main()
